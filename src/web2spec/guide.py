from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image

from .i18n import get_text
from .models import GuideSection, PageSnapshot, SemanticElement
from .utils import ensure_dir, safe_filename_from_url


# Color scheme (RGB)
COLOR_PART_BG = RGBColor(0, 179, 179)  # Teal
COLOR_SECTION_BG = RGBColor(128, 128, 128)  # Grey
COLOR_STEP_BG = RGBColor(144, 238, 144)  # Light green
COLOR_PART_TEXT = RGBColor(255, 255, 255)  # White
COLOR_SECTION_TEXT = RGBColor(255, 255, 255)  # White
COLOR_STEP_TEXT = RGBColor(0, 0, 0)  # Black


def _set_paragraph_shading(paragraph, fill_hex: str) -> None:
    """Apply background shading to a paragraph using docx XML primitives."""
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)


def _add_part_heading(doc: Document, text: str) -> None:
    """Add a part-level heading (CZĘŚĆ) with teal background."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    
    run = p.add_run(text)
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = COLOR_PART_TEXT

    _set_paragraph_shading(p, "00B3B3")  # Teal
    
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_section_heading(doc: Document, text: str, depth: int) -> None:
    """Add a section-level heading (I. Zasady) with grey background."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(10)
    
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = COLOR_SECTION_TEXT

    _set_paragraph_shading(p, "808080")  # Grey
    
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _add_step_heading(doc: Document, step_number: int, heading: str) -> None:
    """Add a step-level heading (1. Logowanie) with light green background."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    
    run = p.add_run(f"{step_number}. {heading}")
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = COLOR_STEP_TEXT

    _set_paragraph_shading(p, "90EE90")  # Light green
    
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _add_step_content(
    doc: Document,
    action_bullets: list[str],
    what_you_see: str,
    screenshot_path: Path | None,
    locale: str = "pl",
) -> None:
    """Add the content for a single step: actions, description, and screenshot."""
    text = get_text(locale)["guide"]
    
    # Action bullets
    if action_bullets:
        p = doc.add_paragraph(text["actions"], style="List Bullet")
        for bullet in action_bullets:
            doc.add_paragraph(bullet, style="List Bullet")
    
    # What you see
    if what_you_see:
        p = doc.add_paragraph()
        run = p.add_run(f"{text['what_you_see']}: ")
        run.bold = True
        run.font.size = Pt(10)
        p.add_run(what_you_see)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
    
    # Screenshot
    if screenshot_path and screenshot_path.exists():
        try:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
            
            # Add image with max width of 6 inches to fit on page
            run = p.add_run()
            run.add_picture(str(screenshot_path), width=Inches(6))
            
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            doc.add_paragraph(f"[Error embedding screenshot: {e}]")


def build_guide(root_url: str, sections: list[GuideSection], locale: str = "pl") -> Document:
    """Build a DOCX document from guide sections."""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    text = get_text(locale)["guide"]
    
    # Title
    title = doc.add_heading(text["guide_title"], 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle with URL
    url_para = doc.add_paragraph(f"URL: {root_url}")
    url_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    url_para.paragraph_format.space_after = Pt(12)
    
    # Content: one part for all sections (simplified; parts could be depth-based if desired)
    _add_part_heading(doc, "Przewodnik")
    
    for section in sections:
        # Section heading (I. Section Title)
        section_heading = f"{_roman_numeral(section.depth + 1)}. {section.title}"
        _add_section_heading(doc, section_heading, section.depth)
        
        # Introduction
        if section.intro:
            intro_para = doc.add_paragraph(section.intro)
            intro_para.paragraph_format.space_after = Pt(10)
        
        # Steps
        for step in section.steps:
            _add_step_heading(doc, step.step_number, step.heading)
            _add_step_content(
                doc,
                step.action_bullets,
                step.what_you_see,
                step.screenshot_path,
                locale,
            )
            
            # Add spacing between steps
            doc.add_paragraph()
    
    return doc


def _roman_numeral(num: int) -> str:
    """Convert integer to Roman numeral."""
    val = [1000, 900, 500, 400, 100, 90, 50, 40, 10, 9, 5, 4, 1]
    syms = ["M", "CM", "D", "CD", "C", "XC", "L", "XL", "X", "IX", "V", "IV", "I"]
    roman_num = ""
    i = 0
    while num > 0:
        for _ in range(num // val[i]):
            roman_num += syms[i]
            num -= val[i]
        i += 1
    return roman_num


def write_guide(path: Path, root_url: str, sections: list[GuideSection], locale: str = "pl") -> None:
    """Render guide sections to a DOCX file."""
    doc = build_guide(root_url, sections, locale)
    doc.save(str(path))


def attach_focused_step_images(
    section: GuideSection,
    snapshot: PageSnapshot,
    crops_dir: Path,
    top_padding: int,
    bottom_padding: int,
) -> GuideSection:
    """Generate focused crops for each guide step when the step references a visible UI control."""
    if snapshot.screenshot_path is None or not snapshot.screenshot_path.exists():
        return section

    ensure_dir(crops_dir)

    for step in section.steps:
        label_candidates = _extract_bracket_labels(step.action_bullets)
        best_match = _find_best_element_match(snapshot.elements, label_candidates)
        if best_match is None:
            step.screenshot_path = snapshot.screenshot_path
            continue

        crop_path = crops_dir / f"{safe_filename_from_url(snapshot.url)}-step-{step.step_number:02d}.png"
        created = _crop_full_width_window(
            snapshot.screenshot_path,
            crop_path,
            int(best_match.bbox.y),
            int(best_match.bbox.height),
            top_padding,
            bottom_padding,
        )
        step.screenshot_path = crop_path if created else snapshot.screenshot_path

    return section


def _extract_bracket_labels(action_bullets: list[str]) -> list[str]:
    labels: list[str] = []
    for bullet in action_bullets:
        labels.extend(match.strip() for match in re.findall(r"\[([^\]]+)\]", bullet) if match.strip())
    return labels


def _normalize_label(value: str) -> str:
    return re.sub(r"\s+", " ", value).strip().casefold()


def _find_best_element_match(elements: list[SemanticElement], labels: list[str]) -> SemanticElement | None:
    if not labels:
        return None

    normalized_labels = [_normalize_label(item) for item in labels if item.strip()]
    best_score = 0
    best_element: SemanticElement | None = None

    for element in elements:
        if element.bbox is None:
            continue
        candidate_texts = [
            element.text,
            element.aria_label,
            element.placeholder,
            element.name,
            element.element_id,
        ]
        normalized_candidates = [_normalize_label(text) for text in candidate_texts if text]
        if not normalized_candidates:
            continue

        score = 0
        for label in normalized_labels:
            for candidate in normalized_candidates:
                if label == candidate:
                    score = max(score, 100)
                elif label in candidate or candidate in label:
                    score = max(score, 60)
                else:
                    overlap = len(set(label.split()) & set(candidate.split()))
                    score = max(score, overlap * 10)

        if score > best_score:
            best_score = score
            best_element = element

    return best_element if best_score >= 20 else None


def _crop_full_width_window(
    source_path: Path,
    target_path: Path,
    bbox_y: int,
    bbox_height: int,
    top_padding: int,
    bottom_padding: int,
) -> bool:
    try:
        with Image.open(source_path) as img:
            width, height = img.size
            top = max(0, bbox_y - top_padding)
            bottom = min(height, bbox_y + max(1, bbox_height) + bottom_padding)
            if bottom <= top:
                return False
            cropped = img.crop((0, top, width, bottom))
            cropped.save(target_path)
            return True
    except Exception:
        return False
